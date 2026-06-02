#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""遍历 Word 文档中的所有测试用例表格，调用 DeepSeek API 重新编写每个测试
步骤的“实际结果”。

设计约束（与需求一一对应）：
  1. 实际结果依据该步骤的“操作过程及数据”和“预期结果”生成；
  2. 结果符合系统测试要求，简洁明了；
  3. 仅修改“实际结果”单元格，不改动用例其它任何内容；
  4. 不改变表格结构与样式（保留原单元格的段落与 run 格式）；
  5. 提示词中将大模型角色设定为资深测试专家；
  6. 调用时把整个用例内容作为参考一并提供给大模型。

用法：
    export DEEPSEEK_API_KEY="sk-xxxx"
    python3 fill_actual_results.py A.docx
    python3 fill_actual_results.py A.docx -o A_filled.docx
    python3 fill_actual_results.py A.docx --dry-run   # 不调用 API，仅打印将要发送的内容
"""

import argparse
import json
import os
import sys
import time

import requests
from docx import Document
from docx.oxml.ns import qn

DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"
DEEPSEEK_MODEL = "deepseek-chat"

# 表头中各列的关键字（用于动态定位列，避免依赖固定列号）
COL_KEYS = {
    "seq": "序号",
    "operation": "操作过程及数据",
    "expected": "预期结果",
    "actual": "实际结果",
    "remark": "备注",
}


def get_cell_text(cell):
    """读取单元格文本（多段落用换行拼接）。"""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def unique_cell_indices(row):
    """返回该行中去重后的 (列索引, 单元格) 列表。

    Word 的横向合并单元格在 python-docx 中会重复返回同一个 tc，
    这里依据底层 tc 元素去重，得到逻辑上的列。
    """
    result = []
    seen = set()
    for idx, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        result.append((idx, cell))
    return result


def find_header(table):
    """在表格中查找测试步骤的表头行，返回 (表头行索引, {字段: 列索引})。

    找不到则返回 (None, None)。
    """
    for ri, row in enumerate(table.rows):
        texts = [get_cell_text(c) for c in row.cells]
        if COL_KEYS["seq"] in texts and COL_KEYS["actual"] in texts:
            col_map = {}
            for key, kw in COL_KEYS.items():
                for ci, cell in unique_cell_indices(row):
                    if get_cell_text(cell) == kw:
                        col_map[key] = ci
                        break
            if "actual" in col_map and "operation" in col_map and "expected" in col_map:
                return ri, col_map
    return None, None


def build_case_context(table):
    """将整个用例表格内容序列化为文本，供大模型参考。"""
    lines = []
    for row in table.rows:
        cells = [get_cell_text(c) for c in unique_cell_indices_cells(row)]
        cells = [c for c in cells if c != ""]
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def unique_cell_indices_cells(row):
    return [cell for _, cell in unique_cell_indices(row)]


def is_step_row(seq_text):
    """判断是否为有效的步骤行（序号为纯数字）。"""
    return seq_text.strip().isdigit()


def set_cell_text_preserve_format(cell, text):
    """将文本写入单元格，并尽量保留原有段落/run 的格式。

    仅保留第一个段落与其第一个 run 的格式，多余段落、run 会被移除，
    从而不改变字体、字号等样式（满足“不改变表结构样式”的要求）。
    支持多行文本（以同一 run 的格式插入换行）。
    """
    paragraphs = cell.paragraphs
    first_p = paragraphs[0]

    # 删除多余段落
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    runs = first_p.runs
    if runs:
        target_run = runs[0]
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        target_run = first_p.add_run("")

    _set_run_text_with_breaks(target_run, text)


def _set_run_text_with_breaks(run, text):
    """设置 run 文本，保留 run 的 rPr（格式），并支持换行。"""
    rpr = run._element.find(qn("w:rPr"))
    # 清空 run 内除 rPr 外的内容
    for child in list(run._element):
        if child is rpr:
            continue
        run._element.remove(child)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br = run._element.makeelement(qn("w:br"), {})
            run._element.append(br)
        t = run._element.makeelement(qn("w:t"), {})
        t.set(qn("xml:space"), "preserve")
        t.text = line
        run._element.append(t)


def build_messages(case_context, operation, expected):
    """构造 DeepSeek Chat 接口的 messages。"""
    system_prompt = (
        "你是一名资深测试专家，精通软件系统测试与测试用例评审。"
        "你的任务是为测试用例中的某个测试步骤编写规范的“实际结果”。"
        "要求：\n"
        "1. 实际结果必须严格依据该步骤的“操作过程及数据”和“预期结果”来编写；\n"
        "2. 实际结果需符合系统测试对测试结果记录的要求，客观、准确、简洁明了；\n"
        "3. 描述实际执行后系统真实呈现的现象，与预期结果相对应；\n"
        "4. 只输出“实际结果”正文，不要输出任何解释、前后缀、标题或引号。"
    )
    user_prompt = (
        "下面是完整的测试用例内容，供你理解上下文：\n"
        "------ 用例内容开始 ------\n"
        f"{case_context}\n"
        "------ 用例内容结束 ------\n\n"
        "请针对以下测试步骤，编写其“实际结果”：\n"
        f"【操作过程及数据】{operation}\n"
        f"【预期结果】{expected}\n\n"
        "请直接输出该步骤简洁明了的实际结果："
    )
    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]


def call_deepseek(api_key, messages, max_retries=4, timeout=60):
    """调用 DeepSeek Chat 接口，带指数退避重试。"""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": messages,
        "temperature": 0.3,
        "stream": False,
    }
    last_err = None
    for attempt in range(max_retries):
        try:
            resp = requests.post(
                DEEPSEEK_API_URL, headers=headers, json=payload, timeout=timeout
            )
            resp.raise_for_status()
            data = resp.json()
            return data["choices"][0]["message"]["content"].strip()
        except Exception as e:  # noqa: BLE001
            last_err = e
            wait = 4 * (2 ** attempt)
            print(f"  [警告] 第 {attempt + 1} 次调用失败：{e}，{wait}s 后重试……")
            time.sleep(wait)
    raise RuntimeError(f"DeepSeek 接口调用失败：{last_err}")


def process_document(doc_path, output_path, api_key, dry_run=False):
    doc = Document(doc_path)
    total_tables = len(doc.tables)
    print(f"文档共包含 {total_tables} 个表格。")

    step_count = 0
    for ti, table in enumerate(doc.tables):
        header_idx, col_map = find_header(table)
        if header_idx is None:
            print(f"表格 {ti}: 未找到测试步骤表头，跳过。")
            continue

        case_context = build_case_context(table)
        actual_col = col_map["actual"]
        op_col = col_map["operation"]
        exp_col = col_map["expected"]
        seq_col = col_map.get("seq", 0)

        print(f"\n===== 表格 {ti}（表头行 {header_idx}）开始处理 =====")
        for ri in range(header_idx + 1, len(table.rows)):
            row = table.rows[ri]
            seq_text = get_cell_text(row.cells[seq_col])
            if not is_step_row(seq_text):
                continue

            operation = get_cell_text(row.cells[op_col])
            expected = get_cell_text(row.cells[exp_col])
            if not operation and not expected:
                continue

            step_count += 1
            print(f"\n步骤 {seq_text}：")
            print(f"  操作过程及数据：{operation}")
            print(f"  预期结果：{expected}")

            if dry_run:
                print("  [dry-run] 跳过 API 调用，不修改文档。")
                continue

            messages = build_messages(case_context, operation, expected)
            actual_result = call_deepseek(api_key, messages)
            print(f"  生成的实际结果：{actual_result}")

            set_cell_text_preserve_format(row.cells[actual_col], actual_result)

    if dry_run:
        print(f"\n[dry-run] 共识别 {step_count} 个测试步骤，未写入文件。")
        return

    doc.save(output_path)
    print(f"\n处理完成，共更新 {step_count} 个步骤的实际结果。已保存至：{output_path}")


def parse_args():
    parser = argparse.ArgumentParser(
        description="调用 DeepSeek API 重写 Word 测试用例中每个步骤的“实际结果”。"
    )
    parser.add_argument("docx", help="输入的 Word 文档路径，例如 A.docx")
    parser.add_argument(
        "-o",
        "--output",
        default=None,
        help="输出文档路径，默认覆盖原文件（建议另存）。",
    )
    parser.add_argument(
        "--api-key",
        default=None,
        help="DeepSeek API Key，缺省时读取环境变量 DEEPSEEK_API_KEY。",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="仅解析并打印每个步骤内容，不调用 API、不修改文档。",
    )
    return parser.parse_args()


def main():
    args = parse_args()

    if not os.path.isfile(args.docx):
        print(f"错误：找不到文件 {args.docx}", file=sys.stderr)
        sys.exit(1)

    output_path = args.output or args.docx

    api_key = args.api_key or os.environ.get("DEEPSEEK_API_KEY")
    if not args.dry_run and not api_key:
        print(
            "错误：未提供 DeepSeek API Key。请通过 --api-key 传入，"
            "或设置环境变量 DEEPSEEK_API_KEY。",
            file=sys.stderr,
        )
        sys.exit(1)

    process_document(args.docx, output_path, api_key, dry_run=args.dry_run)


if __name__ == "__main__":
    main()
