#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""根据「问题.csv」逐条生成「问题报告」和「问题处理报告」。

设计要点（与需求一一对应）：

1. 字段一一对照：CSV 表头与 Word 模板表格中的字段名称完全对应，按名称定位
   模板里对应字段的「值单元格」（即字段名右侧紧邻的单元格）后写入。
2. 不修改 Word 模板表格格式：脚本以模板文件为基础另存为新文件，仅向空白的
   「值单元格」内写入文本，复用单元格原有段落，并沿用同行字段名单元格的字体
   run 格式（字体、字号等）。不增删行列、不改变合并关系、不调整列宽/样式。
3. 模板中已固定的字段（问题级别=一般、测试类型=动态测试）保持不变。

用法：

    # 默认：读取 问题.csv，使用 问题报告.docx / 问题处理报告.docx 模板，
    # 在 输出报告/ 目录下逐条生成两份报告
    python3 generate_reports.py

    # 自定义
    python3 generate_reports.py \
        --csv 问题.csv \
        --problem-template 问题报告.docx \
        --handle-template 问题处理报告.docx \
        --outdir 输出报告
"""

import argparse
import copy
import csv
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell


# ---------------------------------------------------------------------------
# 字段映射：模板表格中的「字段名」-> 该字段在 CSV 中的列名
# 「问题级别」「测试类型」为模板固定值，CSV 中没有，故不在此映射中，保持原样。
# ---------------------------------------------------------------------------
PROBLEM_REPORT_FIELDS = {
    "问题编号": "问题编号",
    "问题描述": "问题描述",
    "测试用例名称": "测试用例名称",
    "模块名称": "模块名称",
    "重现步骤": "重现步骤",
    "测试人员": "测试人员",
    "测试日期": "测试日期",
}

HANDLE_REPORT_FIELDS = {
    "问题编号": "问题编号",
    "问题描述": "问题描述",
    "测试用例名称": "测试用例名称",
    "模块名称": "模块名称",
    "处理措施和结果": "处理措施和结果",
    "测试人员": "测试人员",
    "测试日期": "测试日期",
}


def read_csv_records(path):
    """读取 CSV（自动尝试 GBK / UTF-8 编码），返回 dict 列表。"""
    last_err = None
    for enc in ("gbk", "gb18030", "utf-8-sig", "utf-8"):
        try:
            with open(path, encoding=enc, newline="") as f:
                reader = csv.DictReader(f)
                records = [
                    {(k.strip() if k else k): (v.strip() if isinstance(v, str) else v)
                     for k, v in row.items()}
                    for row in reader
                ]
            return records, enc
        except UnicodeDecodeError as e:  # 编码不对，换下一个
            last_err = e
            continue
    raise RuntimeError("无法识别 CSV 文件编码：%s" % last_err)


def _tc_text(tc):
    """读取一个 w:tc 单元格的纯文本。"""
    return "".join(node.text or "" for node in tc.iter(qn("w:t")))


def _label_run_pr(label_tc):
    """取字段名单元格中第一个 run 的 rPr（用于让值文本沿用同样的字体格式）。"""
    for r in label_tc.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
        return None
    return None


def _set_value_cell(value_tc, text, table_obj, ref_rpr):
    """向「值单元格」写入文本，保留单元格原有段落格式，沿用字段名的 run 格式。

    - 多行（含换行符）按软换行写入到同一段落，避免新增段落破坏行高/格式；
    - 仅替换文本内容，不触碰单元格属性（tcPr）、不改变表结构。
    """
    cell = _Cell(value_tc, table_obj)
    para = cell.paragraphs[0]

    # 清空该段落已有的 run（模板值单元格本就为空，这里防御性清理）
    for r in list(para.runs):
        r._element.getparent().remove(r._element)

    lines = re.split(r"\r\n|\r|\n", text if text is not None else "")
    run = para.add_run(lines[0])
    if ref_rpr is not None:
        run._element.insert(0, copy.deepcopy(ref_rpr))
    for extra in lines[1:]:
        run.add_break()
        run.add_text(extra)


def fill_template(template_path, field_map, record, out_path):
    """基于模板生成一份报告：按字段名定位值单元格并填写。"""
    doc = Document(template_path)
    if not doc.tables:
        raise RuntimeError("模板 %s 中没有找到表格" % template_path)

    filled = set()
    for table in doc.tables:
        for row in table.rows:
            tcs = row._tr.findall(qn("w:tc"))
            for idx, tc in enumerate(tcs):
                label = _tc_text(tc).strip()
                if label in field_map and idx + 1 < len(tcs):
                    value_tc = tcs[idx + 1]
                    ref_rpr = _label_run_pr(tc)
                    value = record.get(field_map[label], "")
                    _set_value_cell(value_tc, value or "", table, ref_rpr)
                    filled.add(label)

    missing = set(field_map) - filled
    if missing:
        print("    [警告] 模板 %s 中未找到字段：%s"
              % (os.path.basename(template_path), "、".join(sorted(missing))))

    doc.save(out_path)


def safe_name(text):
    """把问题编号等转成可用作文件名的安全字符串。"""
    return re.sub(r'[\\/:*?"<>|]+', "_", (text or "").strip()) or "未命名"


def main(argv=None):
    parser = argparse.ArgumentParser(
        description="根据「问题.csv」逐条生成「问题报告」与「问题处理报告」")
    parser.add_argument("--csv", default="问题.csv", help="问题记录 CSV 文件路径")
    parser.add_argument("--problem-template", default="问题报告.docx",
                        help="问题报告模板 docx")
    parser.add_argument("--handle-template", default="问题处理报告.docx",
                        help="问题处理报告模板 docx")
    parser.add_argument("--outdir", default="输出报告", help="输出目录")
    args = parser.parse_args(argv)

    for p in (args.csv, args.problem_template, args.handle_template):
        if not os.path.exists(p):
            print("[错误] 文件不存在：%s" % p, file=sys.stderr)
            return 1

    records, enc = read_csv_records(args.csv)
    print("已读取 %s（编码 %s），共 %d 条记录。" % (args.csv, enc, len(records)))
    if not records:
        print("CSV 中没有数据记录。")
        return 0

    os.makedirs(args.outdir, exist_ok=True)

    for i, record in enumerate(records, 1):
        no = safe_name(record.get("问题编号", "第%d条" % i))
        print("[%d/%d] 生成问题编号 %s 的报告……" % (i, len(records), no))

        problem_out = os.path.join(args.outdir, "问题报告_%s.docx" % no)
        handle_out = os.path.join(args.outdir, "问题处理报告_%s.docx" % no)

        fill_template(args.problem_template, PROBLEM_REPORT_FIELDS, record, problem_out)
        fill_template(args.handle_template, HANDLE_REPORT_FIELDS, record, handle_out)

        print("    -> %s" % problem_out)
        print("    -> %s" % handle_out)

    print("全部完成，输出目录：%s" % os.path.abspath(args.outdir))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
