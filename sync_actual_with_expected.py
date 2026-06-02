#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""遍历 Word 测试用例文档中“指定页面”的所有用例表格，逐个检查每个操作步骤：
当该步骤的“实际结果”内容为“与预期结果一致”时，把同一行“预期结果”的内容
复制到“实际结果”中。

设计约束（与需求一一对应）：
  1. 仅修改满足条件的步骤行中“实际结果”单元格，绝不修改用例标题/编号/描述/
     前置条件/模块名称/测试类型等任何标识信息；
  2. 不修改测试用例的备注、测试结论、测试人员、日期及其它任何内容；
  3. 不改变表格的结构与样式：行列数量、合并关系保持不变，且写入“实际结果”时
     沿用该单元格原有段落/run 的格式（字体、字号等），只替换文本内容。

关于“指定页面”：
  .docx 本身并不存储分页信息（页码由 Word 在排版渲染时实时计算）。本脚本依据
  Word 在最近一次渲染时写入文档的分页标记（w:lastRenderedPageBreak）以及手动
  分页符（w:br w:type="page"）来推算每个表格所在的页码。
    * 若文档中存在上述分页标记，可用 --pages 选择页码（如 "1-3,5"）；
    * 若文档中不存在任何分页标记（例如从未在 Word 中渲染保存过），则无法可靠地
      区分页面，此时脚本默认处理全部用例表格，并给出提示；也可用 --tables 按
      表格序号（从 1 开始）精确选择需要处理的用例表格。

用法：
    # 处理全部用例表格（默认输出到 *_同步.docx）
    python3 sync_actual_with_expected.py 系统测试用例.docx

    # 仅处理第 1~3 页和第 5 页的用例表格
    python3 sync_actual_with_expected.py 系统测试用例.docx --pages 1-3,5

    # 按表格序号选择（第 1、2、4 个用例表格）
    python3 sync_actual_with_expected.py 系统测试用例.docx --tables 1,2,4

    # 指定输出文件
    python3 sync_actual_with_expected.py 系统测试用例.docx -o 系统测试用例_同步.docx

    # 仅解析并打印将要修改的步骤，不写文件
    python3 sync_actual_with_expected.py 系统测试用例.docx --dry-run
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn

# 表头中各列的关键字（用于动态定位列，避免依赖固定列号）
COL_KEYS = {
    "seq": "序号",
    "operation": "操作过程及数据",
    "expected": "预期结果",
    "actual": "实际结果",
    "remark": "备注",
}

# “实际结果”等于该文本时，才把“预期结果”复制过来
DEFAULT_MATCH_TEXT = "与预期结果一致"


def get_cell_text(cell):
    """读取单元格文本（多段落以换行拼接）。"""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def unique_cell_indices(row):
    """返回该行中去重后的 (列索引, 单元格) 列表。

    Word 的横向合并单元格在 python-docx 中会重复返回同一个 tc，
    这里依据底层 tc 元素去重，得到逻辑上的列。
    """
    result = []
    seen = set()
    for idx, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        result.append((idx, cell))
    return result


def find_header(table):
    """在表格中查找测试步骤的表头行，返回 (表头行索引, {字段: 列索引})。

    找不到则返回 (None, None)。
    """
    for ri, row in enumerate(table.rows):
        texts = [get_cell_text(c) for c in row.cells]
        if COL_KEYS["seq"] in texts and COL_KEYS["actual"] in texts:
            col_map = {}
            for key, kw in COL_KEYS.items():
                for ci, cell in unique_cell_indices(row):
                    if get_cell_text(cell) == kw:
                        col_map[key] = ci
                        break
            if "actual" in col_map and "expected" in col_map:
                return ri, col_map
    return None, None


def is_step_row(seq_text):
    """判断是否为有效的步骤行（序号为纯数字）。"""
    return seq_text.strip().isdigit()


def normalize(text):
    """规范化文本以便比较：去除所有空白字符与末尾的中英文句号。"""
    t = re.sub(r"\s+", "", text or "")
    return t.rstrip("。.")


def set_cell_text_preserve_format(cell, text):
    """将文本写入单元格，并尽量保留原有段落/run 的格式。

    仅保留第一个段落与其第一个 run 的格式，多余段落、run 会被移除，
    从而不改变字体、字号等样式（满足“不改变表结构样式”的要求）。
    支持多行文本（以同一 run 的格式插入换行）。
    """
    paragraphs = cell.paragraphs
    first_p = paragraphs[0]

    # 删除多余段落
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    runs = first_p.runs
    if runs:
        target_run = runs[0]
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        target_run = first_p.add_run("")

    _set_run_text_with_breaks(target_run, text)


def _set_run_text_with_breaks(run, text):
    """设置 run 文本，保留 run 的 rPr（格式），并支持换行。"""
    rpr = run._element.find(qn("w:rPr"))
    for child in list(run._element):
        if child is rpr:
            continue
        run._element.remove(child)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br = run._element.makeelement(qn("w:br"), {})
            run._element.append(br)
        t = run._element.makeelement(qn("w:t"), {})
        t.set(qn("xml:space"), "preserve")
        t.text = line
        run._element.append(t)


def count_page_breaks(element):
    """统计某个 body 子元素内部触发分页的标记数量。

    包含 Word 渲染时写入的 w:lastRenderedPageBreak 以及手动分页符
    w:br[@w:type='page']。
    """
    count = 0
    count += len(element.findall(".//" + qn("w:lastRenderedPageBreak")))
    for br in element.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            count += 1
    return count


def compute_table_pages(doc):
    """按文档顺序推算每个表格所在的页码。

    返回 (table_pages, has_page_info)：
      * table_pages：列表，与 doc.tables 一一对应，元素为该表起始处的页码（从 1 起）；
      * has_page_info：文档中是否存在任何分页标记。
    """
    body = doc.element.body
    tbl_elems = [t._tbl for t in doc.tables]
    tbl_index = {id(e): i for i, e in enumerate(tbl_elems)}

    table_pages = [1] * len(tbl_elems)
    page = 1
    total_breaks = 0
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "tbl":
            idx = tbl_index.get(id(child))
            if idx is not None:
                table_pages[idx] = page
            inner = count_page_breaks(child)
            page += inner
            total_breaks += inner
        elif tag == "p":
            inner = count_page_breaks(child)
            page += inner
            total_breaks += inner
    return table_pages, total_breaks > 0


def parse_int_ranges(spec):
    """解析形如 "1-3,5,8" 的选择表达式，返回整数集合。"""
    result = set()
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            a, b = int(a), int(b)
            if a > b:
                a, b = b, a
            result.update(range(a, b + 1))
        else:
            result.add(int(part))
    return result


def select_tables(doc, pages_spec, tables_spec):
    """根据 --pages / --tables 计算需要处理的表格序号集合（0 基）。

    返回 (selected_set 或 None 表示全部, 提示信息列表)。
    """
    notes = []
    table_pages, has_page_info = compute_table_pages(doc)

    if tables_spec:
        wanted = parse_int_ranges(tables_spec)
        selected = {i for i in range(len(doc.tables)) if (i + 1) in wanted}
        notes.append(f"按表格序号选择：{sorted(n + 1 for n in selected)}")
        return selected, notes

    if pages_spec:
        if not has_page_info:
            notes.append(
                "警告：文档中未发现任何分页标记，无法按页码筛选。"
                "将改为处理全部用例表格。如需精确选择，请改用 --tables 指定表格序号。"
            )
            return None, notes
        wanted_pages = parse_int_ranges(pages_spec)
        selected = {
            i for i, pg in enumerate(table_pages) if pg in wanted_pages
        }
        mapping = ", ".join(
            f"表格{ i + 1 }→第{pg}页" for i, pg in enumerate(table_pages)
        )
        notes.append(f"分页推算：{mapping}")
        notes.append(
            f"按页码 {sorted(wanted_pages)} 选择到表格："
            f"{sorted(n + 1 for n in selected)}"
        )
        return selected, notes

    # 未指定任何筛选：处理全部
    if not has_page_info:
        notes.append("提示：未指定 --pages/--tables，将处理全部用例表格。")
    return None, notes


def process_document(doc_path, output_path, pages_spec, tables_spec,
                     match_text, dry_run=False):
    doc = Document(doc_path)
    total_tables = len(doc.tables)
    print(f"文档共包含 {total_tables} 个表格。")

    selected, notes = select_tables(doc, pages_spec, tables_spec)
    for n in notes:
        print(n)

    target_norm = normalize(match_text)
    changed = 0
    scanned_steps = 0
    processed_tables = 0

    for ti, table in enumerate(doc.tables):
        if selected is not None and ti not in selected:
            continue

        header_idx, col_map = find_header(table)
        if header_idx is None:
            print(f"表格 {ti + 1}: 未找到测试步骤表头，跳过。")
            continue

        processed_tables += 1
        actual_col = col_map["actual"]
        exp_col = col_map["expected"]
        seq_col = col_map.get("seq", 0)

        print(f"\n===== 表格 {ti + 1}（表头行 {header_idx}）开始处理 =====")
        for ri in range(header_idx + 1, len(table.rows)):
            row = table.rows[ri]
            seq_text = get_cell_text(row.cells[seq_col])
            if not is_step_row(seq_text):
                continue

            scanned_steps += 1
            actual_cell = row.cells[actual_col]
            actual_text = get_cell_text(actual_cell)
            expected_text = get_cell_text(row.cells[exp_col])

            if normalize(actual_text) != target_norm:
                continue

            print(f"\n步骤 {seq_text}：实际结果=“{actual_text}”")
            print(f"  → 复制预期结果：{expected_text}")
            changed += 1

            if dry_run:
                print("  [dry-run] 未写入文档。")
                continue

            set_cell_text_preserve_format(actual_cell, expected_text)

    print(
        f"\n处理表格 {processed_tables} 个，扫描步骤 {scanned_steps} 个，"
        f"命中并{'将' if dry_run else '已'}复制 {changed} 个步骤的实际结果。"
    )

    if dry_run:
        print("[dry-run] 未写入文件。")
        return

    if changed == 0:
        print("没有需要修改的步骤，未写入文件。")
        return

    doc.save(output_path)
    print(f"已保存至：{output_path}")


def parse_args():
    parser = argparse.ArgumentParser(
        description="将测试用例中“实际结果”为“与预期结果一致”的步骤替换为“预期结果”的内容。"
    )
    parser.add_argument(
        "docx",
        nargs="?",
        default="系统测试用例.docx",
        help="输入的 Word 文档路径，默认 系统测试用例.docx",
    )
    parser.add_argument(
        "-o", "--output", default=None,
        help="输出文档路径，默认在原文件名后追加“_同步”。",
    )
    parser.add_argument(
        "--pages", default=None,
        help='指定页码，如 "1-3,5"。仅当文档含分页标记时有效。',
    )
    parser.add_argument(
        "--tables", default=None,
        help='按用例表格序号（从1起）选择，如 "1,2,4"。优先级高于 --pages。',
    )
    parser.add_argument(
        "--match", default=DEFAULT_MATCH_TEXT,
        help=f"触发复制的“实际结果”文本，默认“{DEFAULT_MATCH_TEXT}”。",
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="仅解析并打印将要修改的步骤，不写入文件。",
    )
    return parser.parse_args()


def default_output_path(docx_path):
    base, ext = os.path.splitext(docx_path)
    return f"{base}_同步{ext}"


def main():
    args = parse_args()

    if not os.path.isfile(args.docx):
        print(f"错误：找不到文件 {args.docx}", file=sys.stderr)
        sys.exit(1)

    output_path = args.output or default_output_path(args.docx)

    process_document(
        args.docx,
        output_path,
        pages_spec=args.pages,
        tables_spec=args.tables,
        match_text=args.match,
        dry_run=args.dry_run,
    )


if __name__ == "__main__":
    main()
