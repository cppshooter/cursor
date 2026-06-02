#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""遍历《系统测试用例.docx》中的所有用例表格，结合《需求规格.docx》中对应功能模块
的“处理流程”，调用 DeepSeek API 站在系统测试角度重写每个测试步骤。

设计约束（与需求一一对应）：
  1. 不修改用例名称、编号、描述、前置条件、模块名称、测试类型等任何标识信息，
     仅重写每一步的“操作过程及数据”“预期结果”“实际结果”三列；
  2. 站在系统测试角度编写步骤与观察结果，拉通各配置项/硬件设备体现功能；
  3. 不修改测试用例的其它任何内容（备注、测试结论、测试人员、日期等均保留）；
  4. 不改变表格结构与样式（步骤行数量保持不变，仅替换单元格文本，保留段落/run 格式）；
  5. 提示词中将大模型角色设定为资深测试专家；
  6. 每编写一个用例时，都把该用例对应功能模块的完整需求（功能描述、处理流程、
     性能需求、异常处理）作为依据提供给大模型；
  7. 每个用例的操作步骤不超过 8 步（保持与原用例一致的步骤数，且不超过 8）。

用法：
    export DEEPSEEK_API_KEY="sk-xxxx"
    python3 optimize_test_cases.py
    python3 optimize_test_cases.py --cases 系统测试用例.docx --spec 需求规格.docx -o 系统测试用例_优化.docx
    python3 optimize_test_cases.py --dry-run   # 不调用 API，仅打印解析结果与提示词
"""

import argparse
import json
import os
import re
import sys
import time

import requests
from docx import Document
from docx.oxml.ns import qn

DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"
DEEPSEEK_MODEL = "deepseek-chat"

MAX_STEPS = 8

# 测试步骤表头各列关键字（按关键字动态定位列，兼容合并单元格导致的列数差异）
COL_KEYS = {
    "seq": "序号",
    "operation": "操作过程及数据",
    "expected": "预期结果",
    "actual": "实际结果",
    "remark": "备注",
}

# 用例元数据标签（label -> 规范字段名）
META_LABELS = [
    "用例标题",
    "用例编号",
    "用例描述",
    "前置条件",
    "模块名称",
    "测试类型",
    "测试结论",
    "测试人员",
    "测试日期",
]

# 从“XXX（SRS_XXX）”中提取模块编码
CODE_RE = re.compile(r"[（(]\s*([A-Za-z][A-Za-z0-9_]+)\s*[）)]")


# --------------------------------------------------------------------------- #
# 通用辅助
# --------------------------------------------------------------------------- #
def get_cell_text(cell):
    """读取单元格文本（多段落用换行拼接）。"""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def unique_logical_cells(row):
    """返回该行去重后的逻辑单元格列表（横向合并的单元格在 python-docx 中会重复返回同一 tc）。"""
    result = []
    seen = set()
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        result.append(cell)
    return result


def unique_cell_indices(row):
    """返回该行去重后的 (列索引, 单元格) 列表。"""
    result = []
    seen = set()
    for idx, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        result.append((idx, cell))
    return result


def extract_code(text):
    """从形如“模拟信号采集功能（SRS_XHMN_XHCJ_MNXHCJ）”的文本中提取编码。"""
    m = CODE_RE.search(text or "")
    return m.group(1) if m else None


# --------------------------------------------------------------------------- #
# 需求规格解析
# --------------------------------------------------------------------------- #
def parse_spec_modules(spec_path):
    """解析《需求规格.docx》，按功能模块（Heading 5）聚合其下所有需求小节。

    返回 dict：module_code -> {
        "code", "name", "处理流程", "full_text"
    }
    其中 full_text 为该模块完整需求（功能描述/处理流程/性能需求/异常处理…）。
    """
    doc = Document(spec_path)
    modules = {}
    cur = None          # 当前模块 dict
    cur_section = None  # 当前小节名（Heading 6 文本）

    def style_name(p):
        try:
            return p.style.name or ""
        except Exception:  # noqa: BLE001
            return ""

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        sname = style_name(p)

        # 功能模块标题（Heading 5）。兼容：标题中带编码且非设备级（Heading 4）。
        if sname == "Heading 5" or (sname not in ("Heading 6",) and "功能" in text and extract_code(text) and sname.startswith("Heading")):
            if sname == "Heading 5":
                code = extract_code(text)
                name = re.sub(r"[（(].*?[）)]", "", text).strip()
                cur = {
                    "code": code,
                    "name": name,
                    "title": text,
                    "sections": [],   # [(section_name, [lines...])]
                }
                cur_section = None
                if code:
                    modules[code] = cur
                else:
                    modules[text] = cur
                continue

        if cur is None:
            continue

        # 小节标题（Heading 6），如 功能描述/处理流程/性能需求/异常处理
        if sname == "Heading 6":
            cur_section = text
            cur["sections"].append((cur_section, []))
            continue

        # 正文，归入当前小节
        if cur["sections"]:
            cur["sections"][-1][1].append(text)
        else:
            cur["sections"].append(("说明", [text]))

    # 生成衍生字段
    for mod in modules.values():
        flow_lines = []
        full_parts = [f"模块：{mod['title']}"]
        for sec_name, lines in mod["sections"]:
            body = "\n".join(lines).strip()
            full_parts.append(f"【{sec_name}】\n{body}")
            if "处理流程" in sec_name:
                flow_lines.append(body)
        mod["处理流程"] = "\n".join(flow_lines).strip()
        mod["full_text"] = "\n\n".join(full_parts).strip()

    return modules


def lookup_module(modules, module_cell_text):
    """根据用例“模块名称”单元格文本，匹配需求模块（先按编码，再按名称）。"""
    code = extract_code(module_cell_text)
    if code and code in modules:
        return modules[code]
    name = re.sub(r"[（(].*?[）)]", "", module_cell_text or "").strip()
    for mod in modules.values():
        if mod["name"] and mod["name"] == name:
            return mod
    # 兜底：名称包含匹配
    for mod in modules.values():
        if name and mod["name"] and (name in mod["name"] or mod["name"] in name):
            return mod
    return None


# --------------------------------------------------------------------------- #
# 用例表格解析
# --------------------------------------------------------------------------- #
def is_step_table(table):
    return find_header(table)[0] is not None


def find_header(table):
    """查找测试步骤表头行，返回 (表头行索引, {字段: 列索引})；找不到返回 (None, None)。"""
    for ri, row in enumerate(table.rows):
        texts = [get_cell_text(c) for c in row.cells]
        if COL_KEYS["seq"] in texts and COL_KEYS["actual"] in texts:
            col_map = {}
            for key, kw in COL_KEYS.items():
                for ci, cell in unique_cell_indices(row):
                    if get_cell_text(cell) == kw:
                        col_map[key] = ci
                        break
            if all(k in col_map for k in ("seq", "operation", "expected", "actual")):
                return ri, col_map
    return None, None


def extract_case_meta(table):
    """从表格中抽取用例元数据（label -> value）。"""
    meta = {}
    for row in table.rows:
        cells = unique_logical_cells(row)
        i = 0
        while i < len(cells):
            label = get_cell_text(cells[i])
            if label in META_LABELS and i + 1 < len(cells):
                value = get_cell_text(cells[i + 1])
                if label not in meta:
                    meta[label] = value
                i += 2
                continue
            i += 1
    return meta


def is_step_seq(seq_text):
    return seq_text.strip().isdigit()


def collect_steps(table, header_idx, col_map):
    """收集“有内容”的步骤行，返回 [{"row_idx", "seq", "operation", "expected", "actual"}]。"""
    steps = []
    seq_col = col_map["seq"]
    op_col = col_map["operation"]
    exp_col = col_map["expected"]
    act_col = col_map["actual"]
    for ri in range(header_idx + 1, len(table.rows)):
        row = table.rows[ri]
        seq_text = get_cell_text(row.cells[seq_col])
        if not is_step_seq(seq_text):
            continue
        operation = get_cell_text(row.cells[op_col])
        expected = get_cell_text(row.cells[exp_col])
        actual = get_cell_text(row.cells[act_col])
        if not operation and not expected and not actual:
            # 空的占位步骤行，保持原样不动
            continue
        steps.append(
            {
                "row_idx": ri,
                "seq": seq_text.strip(),
                "operation": operation,
                "expected": expected,
                "actual": actual,
            }
        )
    return steps


def build_case_context(table):
    """将整个用例表格序列化为文本，供大模型理解上下文。"""
    lines = []
    for row in table.rows:
        cells = [get_cell_text(c) for c in unique_logical_cells(row)]
        cells = [c for c in cells if c != ""]
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# 写回（保留格式）
# --------------------------------------------------------------------------- #
def set_cell_text_preserve_format(cell, text):
    """写入单元格文本并尽量保留首段首 run 的格式，删除多余段落/run，支持换行。"""
    paragraphs = cell.paragraphs
    first_p = paragraphs[0]
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    runs = first_p.runs
    if runs:
        target_run = runs[0]
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        target_run = first_p.add_run("")

    _set_run_text_with_breaks(target_run, text)


def _set_run_text_with_breaks(run, text):
    rpr = run._element.find(qn("w:rPr"))
    for child in list(run._element):
        if child is rpr:
            continue
        run._element.remove(child)
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            run._element.append(run._element.makeelement(qn("w:br"), {}))
        t = run._element.makeelement(qn("w:t"), {})
        t.set(qn("xml:space"), "preserve")
        t.text = line
        run._element.append(t)


# --------------------------------------------------------------------------- #
# 提示词与 DeepSeek 调用
# --------------------------------------------------------------------------- #
SYSTEM_PROMPT = (
    "你是一名资深测试专家，精通系统测试、测试用例设计与评审，熟悉软硬件一体化的"
    "分系统级测试。你擅长依据需求规格中的“处理流程”，从系统测试角度设计可执行、"
    "可观测的测试步骤，并能拉通各配置项与硬件设备来端到端体现被测功能。\n"
    "你的任务：在不改变测试用例标识信息与表格结构的前提下，重写每个测试步骤的"
    "“操作过程及数据”“预期结果”“实际结果”三列内容。\n"
    "严格遵守以下规则：\n"
    "1. 仅依据所提供功能模块的完整需求（尤其是“处理流程”）来优化步骤；\n"
    "2. 站在系统测试角度编写步骤与观察结果，覆盖处理流程的关键环节，拉通相关"
    "配置项/硬件设备（如射频通路、ADC、DDC、存储、频谱分析仪、功率计等），"
    "体现端到端功能闭环；\n"
    "3. 不得修改用例标题、用例编号、用例描述、前置条件、模块名称、测试类型等任何"
    "标识信息，也不得改变步骤的数量；\n"
    "4. 操作过程及数据要具体可执行（含界面操作、配置参数与数据），预期结果要可"
    "判定、可量化，实际结果要与预期结果对应、客观真实；\n"
    f"5. 步骤总数不超过 {MAX_STEPS} 步；\n"
    "6. 严格按要求输出 JSON，不输出任何多余解释、标题或代码块标记。"
)


def build_user_prompt(module, case_meta, case_context, steps):
    n = len(steps)
    existing_steps = []
    for s in steps:
        existing_steps.append(
            {
                "序号": s["seq"],
                "操作过程及数据": s["operation"],
                "预期结果": s["expected"],
                "实际结果": s["actual"],
            }
        )

    module_block = (
        module["full_text"] if module else "（未匹配到对应功能模块需求，请基于用例自身信息谨慎优化）"
    )
    flow_block = (module or {}).get("处理流程", "")

    prompt = (
        "下面提供该测试用例所属功能模块的完整需求，请作为优化测试步骤的唯一依据：\n"
        "====== 功能模块需求（开始）======\n"
        f"{module_block}\n"
        "====== 功能模块需求（结束）======\n\n"
    )
    if flow_block:
        prompt += (
            "其中“处理流程”是设计系统测试步骤的核心依据，重点据此组织步骤：\n"
            f"{flow_block}\n\n"
        )

    prompt += (
        "下面是完整的测试用例内容（用于理解上下文，其中的标识信息不可修改）：\n"
        "------ 用例内容（开始）------\n"
        f"{case_context}\n"
        "------ 用例内容（结束）------\n\n"
        f"该用例当前共有 {n} 个有效测试步骤，原始步骤如下（JSON）：\n"
        f"{json.dumps(existing_steps, ensure_ascii=False, indent=2)}\n\n"
        "请你作为资深测试专家，依据上述功能模块“处理流程”，从系统测试角度重新优化"
        "这些步骤，使其拉通相关配置项与硬件设备、体现端到端功能闭环。要求：\n"
        f"- 必须输出且仅输出 {n} 个步骤，与原步骤数量一致（不得增减步骤，且不超过 {MAX_STEPS} 步）；\n"
        "- 只优化“操作过程及数据”“预期结果”“实际结果”三项，不要输出其它任何字段；\n"
        "- 保持步骤“序号”与原步骤一致；\n"
        "- 操作过程及数据需具体、可执行；预期结果需可判定、可量化；实际结果需与"
        "预期结果一一对应且为合理的真实执行表现；\n"
        "- 严格输出如下 JSON 结构（不要包含其它内容、不要使用代码块标记）：\n"
        '{"steps": [{"序号": "1", "操作过程及数据": "...", "预期结果": "...", "实际结果": "..."}, ...]}'
    )
    return prompt


def call_deepseek(api_key, messages, max_retries=4, timeout=120):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": messages,
        "temperature": 0.3,
        "stream": False,
        "response_format": {"type": "json_object"},
    }
    last_err = None
    for attempt in range(max_retries):
        try:
            resp = requests.post(
                DEEPSEEK_API_URL, headers=headers, json=payload, timeout=timeout
            )
            resp.raise_for_status()
            data = resp.json()
            return data["choices"][0]["message"]["content"].strip()
        except Exception as e:  # noqa: BLE001
            last_err = e
            wait = 4 * (2 ** attempt)
            print(f"    [警告] 第 {attempt + 1} 次调用失败：{e}，{wait}s 后重试……")
            time.sleep(wait)
    raise RuntimeError(f"DeepSeek 接口调用失败：{last_err}")


def parse_steps_response(content):
    """解析模型返回，得到步骤列表 [{操作过程及数据, 预期结果, 实际结果}]。"""
    text = content.strip()
    # 去除可能的代码块标记
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text).strip()
    data = json.loads(text)
    steps = data.get("steps") if isinstance(data, dict) else data
    if not isinstance(steps, list):
        raise ValueError("返回 JSON 中未找到 steps 列表")
    return steps


# --------------------------------------------------------------------------- #
# 主处理流程
# --------------------------------------------------------------------------- #
def process(cases_path, spec_path, output_path, api_key, dry_run=False):
    modules = parse_spec_modules(spec_path)
    print(f"需求规格解析完成，共识别功能模块 {len(modules)} 个：")
    for code, mod in modules.items():
        has_flow = "有" if mod["处理流程"] else "无"
        print(f"  - {code}：{mod['name']}（处理流程：{has_flow}）")

    doc = Document(cases_path)
    print(f"\n测试用例文档共包含 {len(doc.tables)} 个表格。")

    case_count = 0
    updated_steps = 0
    for ti, table in enumerate(doc.tables):
        header_idx, col_map = find_header(table)
        if header_idx is None:
            print(f"\n表格 {ti}: 未识别到测试步骤表头，跳过。")
            continue

        meta = extract_case_meta(table)
        module_cell = meta.get("模块名称", "")
        module = lookup_module(modules, module_cell)
        steps = collect_steps(table, header_idx, col_map)

        case_count += 1
        title = meta.get("用例标题", "")
        code = meta.get("用例编号", "")
        print(f"\n===== 用例 {ti}：{title}（{code}）=====")
        print(f"  模块名称：{module_cell} -> 匹配：{module['code'] if module else '未匹配'}")
        print(f"  有效步骤数：{len(steps)}")

        if not steps:
            print("  无有效步骤，跳过。")
            continue
        if len(steps) > MAX_STEPS:
            print(f"  [警告] 该用例步骤数 {len(steps)} 超过 {MAX_STEPS}，仍保持原结构按 1:1 优化。")

        case_context = build_case_context(table)
        user_prompt = build_user_prompt(module, meta, case_context, steps)

        if dry_run:
            print("  [dry-run] 生成的用户提示词如下（截断显示）：")
            print("  " + "\n  ".join(user_prompt.splitlines()[:6]))
            print(f"  …（提示词总长 {len(user_prompt)} 字符）")
            continue

        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ]
        content = call_deepseek(api_key, messages)
        new_steps = parse_steps_response(content)

        if len(new_steps) != len(steps):
            print(
                f"  [警告] 模型返回步骤数 {len(new_steps)} 与原步骤数 {len(steps)} 不一致，"
                "按最小数量对齐写回，未覆盖到的步骤保持原样。"
            )

        op_col = col_map["operation"]
        exp_col = col_map["expected"]
        act_col = col_map["actual"]
        for s, new in zip(steps, new_steps):
            row = table.rows[s["row_idx"]]
            new_op = str(new.get("操作过程及数据", s["operation"]))
            new_exp = str(new.get("预期结果", s["expected"]))
            new_act = str(new.get("实际结果", s["actual"]))
            set_cell_text_preserve_format(row.cells[op_col], new_op)
            set_cell_text_preserve_format(row.cells[exp_col], new_exp)
            set_cell_text_preserve_format(row.cells[act_col], new_act)
            updated_steps += 1
            print(f"    步骤 {s['seq']} 已优化。")

    if dry_run:
        print(f"\n[dry-run] 共识别 {case_count} 个用例，未调用 API、未写文件。")
        return

    doc.save(output_path)
    print(
        f"\n处理完成：共优化 {case_count} 个用例、{updated_steps} 个步骤。"
        f"已保存至：{output_path}"
    )


def parse_args():
    parser = argparse.ArgumentParser(
        description="结合需求规格的“处理流程”，调用 DeepSeek 重写系统测试用例步骤。"
    )
    parser.add_argument(
        "--cases", default="系统测试用例.docx", help="系统测试用例 Word 文档路径。"
    )
    parser.add_argument(
        "--spec", default="需求规格.docx", help="需求规格 Word 文档路径。"
    )
    parser.add_argument(
        "-o", "--output", default=None,
        help="输出文档路径，默认在原文件名后加 _优化 另存。",
    )
    parser.add_argument(
        "--api-key", default=None,
        help="DeepSeek API Key，缺省时读取环境变量 DEEPSEEK_API_KEY。",
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="仅解析并打印用例与提示词，不调用 API、不修改文档。",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    for path in (args.cases, args.spec):
        if not os.path.isfile(path):
            print(f"错误：找不到文件 {path}", file=sys.stderr)
            sys.exit(1)

    if args.output:
        output_path = args.output
    else:
        base, ext = os.path.splitext(args.cases)
        output_path = f"{base}_优化{ext}"

    api_key = args.api_key or os.environ.get("DEEPSEEK_API_KEY")
    if not args.dry_run and not api_key:
        print(
            "错误：未提供 DeepSeek API Key。请通过 --api-key 传入，"
            "或设置环境变量 DEEPSEEK_API_KEY。",
            file=sys.stderr,
        )
        sys.exit(1)

    process(args.cases, args.spec, output_path, api_key, dry_run=args.dry_run)


if __name__ == "__main__":
    main()
